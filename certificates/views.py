from django.shortcuts import render, get_object_or_404
from django.http import HttpResponse
from .models import Certificate
from .forms import CertificateForm
from docx import Document
import os

def fill_template(template_path, output_path, data):
    doc = Document(template_path)
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, str(value))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, str(value))
    doc.save(output_path)
    return output_path

def index(request):
    if request.method == 'POST':
        form = CertificateForm(request.POST)
        if form.is_valid():
            serial_number = form.cleaned_data['serial_number']
            pin = form.cleaned_data['pin']
            try:
                certificate = Certificate.objects.get(serial_number=serial_number)
                if certificate.pin == pin:
                    data = {
                        '{Serial Number}': certificate.serial_number,
                        '{Article Number}': certificate.article_number,
                        '{Date End}': certificate.date_end,
                        '{Name pump}': certificate.name_pump,
                        '{WC Number}': certificate.wc_number,
                        '{Name partner}': certificate.name_partner,
                    }
                    template_path = 'templates/template.DOCX'
                    output_path = f'generated_certificates/warranty_certificate_{serial_number}.docx'
                    fill_template(template_path, output_path, data)
                    with open(output_path, 'rb') as f:
                        response = HttpResponse(f.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                        response['Content-Disposition'] = f'attachment; filename={os.path.basename(output_path)}'
                        return response
                else:
                    return HttpResponse("PIN помилковий")
            except Certificate.DoesNotExist:
                return HttpResponse("Серійний номер не знайдено")
    else:
        form = CertificateForm()
    return render(request, 'certificates/index.html', {'form': form})
